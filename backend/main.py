from fastapi import FastAPI
from pydantic import BaseModel
from typing import List, Dict
from docx import Document
from fastapi.middleware.cors import CORSMiddleware
from typing import Any

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # allow all (for dev)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMPLATES = {}

# -------------------------
# Models
# -------------------------
class Element(BaseModel):
    id: str
    type: str
    content: Any = ""   # ✅ accept JSON (Slate)
    data: List = []
    rows: int = 0
    cols: int = 0
    rules: List = []

class Template(BaseModel):
    name: str
    elements: List[Element]

# -------------------------
# Helpers
# -------------------------
def replace_vars(text, data):
    if not text:
        return ""
    for k, v in data.items():
        text = text.replace(f"{{{{{k}}}}}", str(v))
    return text

def evaluate_condition(condition: str, data: dict):
    try:
        return eval(condition, {}, data)
    except Exception as e:
        print("Rule error:", e)
        return False

def resolve_text_switch(el, data):
    for rule in el.get("rules", []):
        if evaluate_condition(rule["condition"], data):
            return rule["value"]
    return ""

# -------------------------
# APIs
# -------------------------
@app.post("/save-template")
def save_template(template: Template):
    TEMPLATES[template.name] = template.dict()
    return {"status": "saved", "template": template.name}

@app.get("/get-template/{name}")
def get_template(name: str):
    return TEMPLATES.get(name, {})

@app.post("/generate-docx/{name}")
def generate_docx(name: str, data: Dict):
    template = TEMPLATES.get(name)

    if not template:
        return {"error": "Template not found"}

    doc = Document()

    # HEADER
    for el in template["elements"]:
        if el["type"] == "header":
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].text = replace_vars(el["content"], data)

    # BODY
    for el in template["elements"]:

        if el["type"] == "text":
            doc.add_paragraph(replace_vars(el["content"], data))

        elif el["type"] == "dynamic-text":
            doc.add_paragraph(replace_vars(el["content"], data))

        elif el["type"] == "text-switch":
            content = resolve_text_switch(el, data)
            doc.add_paragraph(content)

        elif el["type"] == "table":
            table = doc.add_table(rows=el["rows"], cols=el["cols"])
            for i, row in enumerate(el["data"]):
                for j, cell in enumerate(row):
                    table.rows[i].cells[j].text = replace_vars(cell, data)

    # FOOTER
    for el in template["elements"]:
        if el["type"] == "footer":
            section = doc.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = replace_vars(el["content"], data)

    file_path = f"{name}.docx"
    doc.save(file_path)

    return {"status": "generated", "file": file_path}